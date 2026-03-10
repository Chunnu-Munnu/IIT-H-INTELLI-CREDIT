"""
Word CAM generator using python-docx.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from loguru import logger
import os
from app.constants import CRORE, LAKH


def format_inr(paise: int) -> str:
    rupees = paise / 100
    if abs(rupees) >= CRORE:
        return f"₹{rupees / CRORE:.2f} Cr"
    elif abs(rupees) >= LAKH:
        return f"₹{rupees / LAKH:.2f} L"
    return f"₹{rupees:,.0f}"


DECISION_COLORS = {
    "APPROVE": RGBColor(0, 128, 0),
    "APPROVE_WITH_CONDITIONS": RGBColor(255, 165, 0),
    "REFER": RGBColor(255, 140, 0),
    "REJECT": RGBColor(220, 0, 0),
}


class WordExporter:

    def export(self, cam_data: dict, output_path: str) -> str:
        doc = Document()

        # ─── Styles ────────────────────────────────────────────
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # ─── Header ────────────────────────────────────────────
        header = doc.add_heading('CREDIT APPRAISAL MEMORANDUM', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.runs[0].font.color.rgb = RGBColor(31, 56, 100)

        subheader = doc.add_paragraph(f"Prepared by Intelli-Credit AI Engine")
        subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subheader.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph(f"Date: {cam_data.get('generated_at', '')}")
        doc.add_paragraph(f"Company: {cam_data.get('company_name', '')}")
        doc.add_paragraph(f"Case ID: {cam_data.get('case_id', '')}")
        doc.add_paragraph("CONFIDENTIAL — For Internal Bank Use Only").bold = True

        doc.add_paragraph()

        # ─── Decision Badge ─────────────────────────────────────
        decision = cam_data.get("decision", "N/A")
        p = doc.add_paragraph()
        run = p.add_run(f"RECOMMENDATION: {decision}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = DECISION_COLORS.get(decision, RGBColor(0, 0, 0))

        # ─── Credit Score Summary ────────────────────────────────
        doc.add_heading("EXECUTIVE SUMMARY", 1)
        summary_table = doc.add_table(rows=1, cols=4)
        summary_table.style = 'Table Grid'
        hdr = summary_table.rows[0].cells
        hdr[0].text = f"Credit Score\n{cam_data.get('credit_score', 'N/A')}/850"
        hdr[1].text = f"Risk Grade\n{cam_data.get('risk_grade', 'N/A')}"
        hdr[2].text = f"Default Probability\n{cam_data.get('default_probability', 0)*100:.1f}%"
        hdr[3].text = f"Recommended Limit\n{format_inr(cam_data.get('recommended_limit_paise', 0))}"

        doc.add_paragraph()

        # ─── Five Cs ─────────────────────────────────────────────
        doc.add_heading("FIVE Cs ASSESSMENT", 1)
        five_cs = cam_data.get("five_cs", {})
        if five_cs:
            five_c_table = doc.add_table(rows=1, cols=3)
            five_c_table.style = 'Table Grid'
            five_c_table.rows[0].cells[0].text = "Dimension"
            five_c_table.rows[0].cells[1].text = "Score (0-100)"
            five_c_table.rows[0].cells[2].text = "Assessment"
            for dim, score in five_cs.items():
                if dim == "Composite":
                    continue
                row = five_c_table.add_row()
                row.cells[0].text = dim
                row.cells[1].text = str(score)
                row.cells[2].text = "Strong" if score >= 70 else ("Adequate" if score >= 50 else "Weak")

        doc.add_paragraph()

        # ─── Reasons ─────────────────────────────────────────────
        doc.add_heading("DECISION RATIONALE", 1)
        for reason in cam_data.get("reasons", []):
            doc.add_paragraph(reason, style='List Bullet')

        # ─── Interest Rate ────────────────────────────────────────
        doc.add_heading("LOAN STRUCTURE", 1)
        rate = cam_data.get("interest_rate_pct", 0)
        doc.add_paragraph(f"Interest Rate: {rate}% p.a. (MCLR + spread)")
        fb = cam_data.get("facility_breakup", {})
        if fb:
            loan_table = doc.add_table(rows=1, cols=2)
            loan_table.style = 'Table Grid'
            loan_table.rows[0].cells[0].text = "Facility"
            loan_table.rows[0].cells[1].text = "Amount"
            for facility, paise in fb.items():
                row = loan_table.add_row()
                row.cells[0].text = facility.replace("_", " ").title()
                row.cells[1].text = format_inr(paise)

        doc.add_paragraph()

        # ─── Covenants ────────────────────────────────────────────
        doc.add_heading("COVENANTS & CONDITIONS", 1)
        for covenant in cam_data.get("covenants", []):
            doc.add_paragraph(covenant, style='List Bullet')

        # ─── AI Explanation ───────────────────────────────────────
        doc.add_heading("AI MODEL EXPLANATION (SHAP)", 1)
        narrative = cam_data.get("score_narrative", "")
        if narrative:
            doc.add_paragraph(narrative)

        # ─── Financial Tables ─────────────────────────────────────
        financial_records = cam_data.get("financial_records", [])
        if financial_records:
            doc.add_heading("FINANCIAL SUMMARY", 1)
            fin_table = doc.add_table(rows=1, cols=4)
            fin_table.style = 'Table Grid'
            fin_table.rows[0].cells[0].text = "Period"
            fin_table.rows[0].cells[1].text = "Revenue (Cr)"
            fin_table.rows[0].cells[2].text = "PAT (Cr)"
            fin_table.rows[0].cells[3].text = "D/E Ratio"
            for rec in financial_records[:3]:
                row = fin_table.add_row()
                period = rec.get("period", {}) if isinstance(rec, dict) else {}
                row.cells[0].text = period.get("fy_label", "N/A")
                rev = rec.get("revenue_from_ops", 0) if isinstance(rec, dict) else 0
                pat = rec.get("pat", 0) if isinstance(rec, dict) else 0
                row.cells[1].text = f"{(rev or 0)/(CRORE*100):.2f}" if rev else "N/A"
                row.cells[2].text = f"{(pat or 0)/(CRORE*100):.2f}" if pat else "N/A"
                row.cells[3].text = "N/A"

        # ─── Footer ───────────────────────────────────────────────
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "CONFIDENTIAL — Prepared by Intelli-Credit AI Engine | For Authorized Bank Personnel Only"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(output_path)
        logger.info(f"CAM Word exported: {output_path}")
        return output_path
